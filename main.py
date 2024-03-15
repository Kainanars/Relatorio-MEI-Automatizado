from datetime import datetime, timedelta

from docx import Document

import locale

locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

arquivos = 0
fileNames = []

def numero_formatado(numero):
    return locale.currency(numero, grouping=True, symbol=None)


def gerar_intervalo_mes(mes_str, ano_str):
    mes = int(mes_str)
    ano = int(ano_str)
    primeiro_dia = datetime(ano, mes, 1)
    ultimo_dia = primeiro_dia.replace(month=mes + 1, day=1) - timedelta(days=1)
    return f"{primeiro_dia.strftime('%d/%m/%Y')} à {ultimo_dia.strftime('%d/%m/%Y')}"


def set_atividade():
    atividade = int(input("Insira a atividade da empresa:"
                          "\n[1] Comércio"
                          "\n[2] Indústria"
                          "\n[3] Serviço"
                          "\n[4] Comércio e Indústria"
                          "\n[5] Comércio e Serviço"
                          "\n[6] Indústria e Serviço"
                          "\n[7] Comércio, Indústria e Serviço\n"))
    return atividade


# def set_valores(atividade):
#     def comercio():
#         print("Insira os valores de Comércio")
#         semnota = input("Sem nota fiscal: ")
#         semnota = float(semnota) if semnota else 0.0
#         nota = input("Com nota fiscal: ")
#         nota = float(nota) if nota else 0.0
#         return semnota, nota
# 
#     def industria():
#         print("Insira os valores de Indústria")
#         semnota = input("Sem nota fiscal: ")
#         semnota = float(semnota) if semnota else 0.0
#         nota = input("Com nota fiscal: ")
#         nota = float(nota) if nota else 0.0
#         return semnota, nota
# 
#     def servico():
#         print("Insira os valores de Serviço")
#         semnota = input("Sem nota fiscal: ")
#         semnota = float(semnota) if semnota else 0.0
#         nota = input("Com nota fiscal: ")
#         nota = float(nota) if nota else 0.0
#         return semnota, nota
# 
#     switch = {
#         1: comercio,
#         2: industria,
#         3: servico,
#         4: lambda: (comercio(), industria()),
#         5: lambda: (comercio(), servico()),
#         6: lambda: (industria(), servico()),
#         7: lambda: (comercio(), industria(), servico())
#     }
# 
#     return atividade, switch.get(atividade, lambda: (0, 0))()
# 

def set_font(document, font_name):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name


def gerar_copias_intervalo(data_inicial, data_final):
    data_atual = datetime.strptime(data_inicial, "%d/%m/%Y")
    data_fim = datetime.strptime(data_final, "%d/%m/%Y")
    data_init = datetime.strptime(data_inicial, "%d/%m/%Y")
    intervalo_datas = []
    while data_atual <= data_fim:
        mes = data_atual.strftime("%m")
        ano = data_atual.strftime("%Y")
        if data_atual.month == data_init.month:
            primeiro_dia = data_atual
        else:
            primeiro_dia = data_atual.replace(day=1)
        if data_atual.month == data_fim.month and data_atual.year == data_fim.year:
            ultimo_dia_mes = data_fim
        elif data_atual.month == data_init.month:
            copy_primeiro_dia = data_atual.replace(day=1)
            ultimo_dia_mes = (copy_primeiro_dia.replace(month=int(mes) + 1) - timedelta(days=1))
        else:
            ultimo_dia_mes = (primeiro_dia.replace(month=int(mes) + 1) - timedelta(days=1))
        intervalo_datas.append(f"{primeiro_dia.strftime('%d/%m/%Y')} à {ultimo_dia_mes.strftime('%d/%m/%Y')}")
        data_atual = primeiro_dia.replace(day=1) + timedelta(days=32)
        if data_atual <= data_fim:
            data_atual = data_atual.replace(day=int(data_inicial.split("/")[0]))

    return intervalo_datas


def gerar_documento(documento, data, nome, cnpj, atividade, valores, id):
    totalComercio = 0.0
    totalIndustria = 0.0
    totalServico = 0.0

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if "AAA" in celula.text:
                    celula.text = celula.text.replace("AAA", cnpj)
                if "BBB" in celula.text:
                    celula.text = celula.text.replace("BBB", nome)
                if "CCC" in celula.text:
                    data = data
                    celula.text = celula.text.replace("CCC", data)

                if atividade == 1:
                    semnota, nota = valores
                    totalComercio = semnota + nota
                    if "DDD" in celula.text:
                        celula.text = celula.text.replace("DDD", f"{numero_formatado(semnota)}")
                    if "EEE" in celula.text:
                        celula.text = celula.text.replace("EEE", f"{numero_formatado(nota)}")
                    if "FFF" in celula.text:
                        celula.text = celula.text.replace("FFF", f"{numero_formatado(totalComercio)}")

                if atividade == 2:
                    semnota, nota = valores
                    totalIndustria = semnota + nota
                    if "GGG" in celula.text:
                        celula.text = celula.text.replace("GGG", f"{numero_formatado(semnota)}")
                    if "HHH" in celula.text:
                        celula.text = celula.text.replace("HHH", f"{numero_formatado(nota)}")
                    if "III" in celula.text:
                        celula.text = celula.text.replace("III", f"{numero_formatado(totalIndustria)}")

                if atividade == 3:
                    semnota, nota = valores
                    totalServico = semnota + nota
                    if "JJJ" in celula.text:
                        celula.text = celula.text.replace("JJJ", f"{numero_formatado(semnota)}")
                    if "KKK" in celula.text:
                        celula.text = celula.text.replace("KKK", f"{numero_formatado(nota)}")
                    if "LLL" in celula.text:
                        celula.text = celula.text.replace("LLL", f"{numero_formatado(totalServico)}")

                if atividade == 4:
                    semnotaComercio, notaComercio = valores[0]
                    semnotaIndustria, notaIndustria = valores[1]
                    totalComercio = semnotaComercio + notaComercio
                    totalIndustria = semnotaIndustria + notaIndustria
                    if "DDD" in celula.text:
                        celula.text = celula.text.replace("DDD", f"{numero_formatado(semnotaComercio)}")
                    if "EEE" in celula.text:
                        celula.text = celula.text.replace("EEE", f"{numero_formatado(notaComercio)}")
                    if "FFF" in celula.text:
                        celula.text = celula.text.replace("FFF", f"{numero_formatado(totalComercio)}")
                    if "GGG" in celula.text:
                        celula.text = celula.text.replace("GGG", f"{numero_formatado(semnotaIndustria)}")
                    if "HHH" in celula.text:
                        celula.text = celula.text.replace("HHH", f"{numero_formatado(notaIndustria)}")
                    if "III" in celula.text:
                        celula.text = celula.text.replace("III", f"{numero_formatado(totalIndustria)}")

                if atividade == 5:
                    semnotaComercio, notaComercio = valores[0]
                    semnotaServico, notaServico = valores[1]
                    totalComercio = semnotaComercio + notaComercio
                    totalServico = semnotaServico + notaServico
                    if "DDD" in celula.text:
                        celula.text = celula.text.replace("DDD", f"{numero_formatado(semnotaComercio)}")
                    if "EEE" in celula.text:
                        celula.text = celula.text.replace("EEE", f"{numero_formatado(notaComercio)}")
                    if "FFF" in celula.text:
                        celula.text = celula.text.replace("FFF", f"{numero_formatado(totalComercio)}")
                    if "JJJ" in celula.text:
                        celula.text = celula.text.replace("JJJ", f"{numero_formatado(semnotaServico)}")
                    if "KKK" in celula.text:
                        celula.text = celula.text.replace("KKK", f"{numero_formatado(notaServico)}")
                    if "LLL" in celula.text:
                        celula.text = celula.text.replace("LLL", f"{numero_formatado(totalServico)}")

                if atividade == 6:
                    semnotaIndustria, notaIndustria = valores[0]
                    semnotaServico, notaServico = valores[1]
                    totalIndustria = semnotaIndustria + notaIndustria
                    totalServico = semnotaServico + notaServico
                    if "GGG" in celula.text:
                        celula.text = celula.text.replace("GGG", f"{numero_formatado(semnotaIndustria)}")
                    if "HHH" in celula.text:
                        celula.text = celula.text.replace("HHH", f"{numero_formatado(notaIndustria)}")
                    if "III" in celula.text:
                        celula.text = celula.text.replace("III", f"{numero_formatado(totalIndustria)}")
                    if "JJJ" in celula.text:
                        celula.text = celula.text.replace("JJJ", f"{numero_formatado(semnotaServico)}")
                    if "KKK" in celula.text:
                        celula.text = celula.text.replace("KKK", f"{numero_formatado(notaServico)}")
                    if "LLL" in celula.text:
                        celula.text = celula.text.replace("LLL", f"{numero_formatado(totalServico)}")

                if atividade == 7:
                    semnotaComercio, notaComercio = valores[0]
                    semnotaIndustria, notaIndustria = valores[1]
                    semnotaServico, notaServico = valores[2]
                    totalComercio = semnotaComercio + notaComercio
                    totalIndustria = semnotaIndustria + notaIndustria
                    totalServico = semnotaServico + notaServico
                    if "DDD" in celula.text:
                        celula.text = celula.text.replace("DDD", f"{numero_formatado(semnotaComercio)}")
                    if "EEE" in celula.text:
                        celula.text = celula.text.replace("EEE", f"{numero_formatado(notaComercio)}")
                    if "FFF" in celula.text:
                        celula.text = celula.text.replace("FFF", f"{numero_formatado(totalComercio)}")
                    if "GGG" in celula.text:
                        celula.text = celula.text.replace("GGG", f"{numero_formatado(semnotaIndustria)}")
                    if "HHH" in celula.text:
                        celula.text = celula.text.replace("HHH", f"{numero_formatado(notaIndustria)}")
                    if "III" in celula.text:
                        celula.text = celula.text.replace("III", f"{numero_formatado(totalIndustria)}")
                    if "JJJ" in celula.text:
                        celula.text = celula.text.replace("JJJ", f"{numero_formatado(semnotaServico)}")
                    if "KKK" in celula.text:
                        celula.text = celula.text.replace("KKK", f"{numero_formatado(notaServico)}")
                    if "LLL" in celula.text:
                        celula.text = celula.text.replace("LLL", f"{numero_formatado(totalServico)}")

                if "MMM" in celula.text:
                    totalGeral = (totalComercio + totalServico + totalIndustria)
                    print(f"Valor total:{totalGeral}")
                    celula.text = celula.text.replace("MMM", f"{numero_formatado(totalGeral)}")

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if "DDD" in celula.text:
                    celula.text = celula.text.replace("DDD", "")
                if "EEE" in celula.text:
                    celula.text = celula.text.replace("EEE", "")
                if "FFF" in celula.text:
                    celula.text = celula.text.replace("FFF", "")
                if "GGG" in celula.text:
                    celula.text = celula.text.replace("GGG", "")
                if "HHH" in celula.text:
                    celula.text = celula.text.replace("HHH", "")
                if "III" in celula.text:
                    celula.text = celula.text.replace("III", "")
                if "JJJ" in celula.text:
                    celula.text = celula.text.replace("JJJ", "")
                if "KKK" in celula.text:
                    celula.text = celula.text.replace("KKK", "")
                if "LLL" in celula.text:
                    celula.text = celula.text.replace("LLL", "")
                if "MMM" in celula.text:
                    celula.text = celula.text.replace("MMM", "")

    set_font(documento, 'Times New Roman')
    novo_document = f"Relatorio {nome} - {id}.docx"
    fileNames.append(novo_document)
    documento.save(f"Relatorio {nome} - {id}.docx")

# # cnpj = input("Digite o CNPJ: ")
# # nome = input("Digite a Razão Social: ")
# # inicio = input("Digite a data de inicio: ")
# # fim = input("Digite a data final: ")
# cnpj = "11.064.624.0001-99"
# nome = "LEIA DO CARMO DA SILVA FILADELFO"
# inicio = "01/01/2023"
# fim = "20/02/2023"
# atividadeselecionada = set_atividade()
# datas = gerar_copias_intervalo(inicio, fim)
# 
# for data in datas:
#     atividade, valores = set_valores(atividadeselecionada)
#     novo_documento = Document("RELATORIO.docx")
#     arquivos += 1
#     gerar_documento(novo_documento, data, nome, cnpj, atividade, valores)
def filenames():
    return fileNames