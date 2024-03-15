from tkinter import *
from tkinter import ttk

import aspose.words as aw
from docx import Document
from tkcalendar import Calendar

from main import gerar_copias_intervalo, gerar_documento, filenames
from valores import set_valores


def saveWord(nome, cnpj):
    fileNames = filenames()
    
    output = aw.Document()
    output.remove_all_children()
    
    for fileName in fileNames:
        input = aw.Document(fileName)
        # Anexe o documento de origem ao final do documento de destino.
        output.append_document(input, aw.ImportFormatMode.KEEP_SOURCE_FORMATTING)
    
    output.save(f"Relatorio {nome} - {cnpj}.docx")


def run(cnpj, nome, atividades):
    inicio = calendarInit.get_date()
    fim = calendarFim.get_date()
    id = 0
    datas = gerar_copias_intervalo(inicio, fim)

    for data in datas:
        atividade, valores = set_valores(atividades)
        novo_documento = Document("RELATORIO.docx")
        id += 1
        gerar_documento(novo_documento, data, nome, cnpj, atividade, valores, id)
    saveWord(nome, cnpj)

janela = Tk()
janela.title("Relatório MEI Automatizado")

style = ttk.Style()
style.theme_use('alt')  # Escolha de um tema (pode ser 'default', 'alt', 'clam', 'classic')

style.configure('TButton', font=('Helvetica', 12), foreground='black', background='lightblue')  # Configuração do estilo do botão
style.configure('TLabel', font=('Helvetica', 12))  # Configuração do estilo do label

cnpj = ttk.Label(janela, text="Digite o CNPJ:")
cnpj.grid(column=0, row=0)
cnpjValue = Entry(janela, width=100)
cnpjValue.grid(column=0, row=1)

nome = ttk.Label(janela, text="Digite a Razão Social:")
nome.grid(column=1, row=0)
nomeValue = Entry(janela, width=100)
nomeValue.grid(column=1, row=1)

inicio = ttk.Label(janela, text="Selecione a data de início:")
inicio.grid(column=0, row=2)

calendarInit = Calendar(janela, date_pattern='dd/mm/yyyy')
calendarInit.grid(column=0, row=3)

fim = ttk.Label(janela, text="Selecione a data de fim:")
fim.grid(column=1, row=2)

calendarFim = Calendar(janela, date_pattern='dd/mm/yyyy')
calendarFim.grid(column=1, row=3)

atividade_frame = ttk.Frame(janela)
atividade_frame.grid(column=0, row=5, columnspan=2)

atividades = []

def toggle_atividade(value):
    if value in atividades:
        atividades.remove(value)
    else:
        atividades.append(value)

comercio_button = ttk.Checkbutton(atividade_frame, text="Comércio", command=lambda: toggle_atividade("Comércio"))
comercio_button.grid(column=0, row=0, padx=5, pady=5)
industria_button = ttk.Checkbutton(atividade_frame, text="Indústria", command=lambda: toggle_atividade("Indústria"))
industria_button.grid(column=1, row=0, padx=5, pady=5)
servico_button = ttk.Checkbutton(atividade_frame, text="Serviço", command=lambda: toggle_atividade("Serviço"))
servico_button.grid(column=2, row=0, padx=5, pady=5)

botao = ttk.Button(janela, text="Avançar", command=lambda: run(cnpjValue.get(), nomeValue.get(), atividades))
botao.grid(column=1, row=5)

janela.mainloop()
